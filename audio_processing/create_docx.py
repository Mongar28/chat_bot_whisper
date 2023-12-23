import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def docx_generator(message, texto, nombre_audio, meta_audio):
# Crea un nuevo documento

    encabezado: str = '''--------------------------------------------------------------------------------------------------------------------
Esta transcripción es producto de los desarrollos del equipo de la Linea
de investigación aplicada en Territorios Inteligentes, que hace parte del
grupo de investigación: Redes y Actores Sociales (RAS) del departamento de
Sociología de la Facultad de Ciencias Sociales y Humanas de la Universidad
de Antioquia. Su efectividad está supeditado a la calidad del audio, téngalo
en cuenta al momento de revisarlo.\t
--------------------------------------------------------------------------------------------------------------------'''

    nombre_archivo = f'documentos/{nombre_audio}.docx'
    doc = docx.Document()
    
    
    # Encabezado
    titulo = doc.add_paragraph(encabezado, style= 'Body Text')
    titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada
    
    # Metadatos
    doc.add_paragraph(meta_audio)
    
    # texto
    paragraph = doc.add_paragraph(texto, style= 'Body Text')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada

    # Guarda el documento en un archivo .docx
    doc.save(nombre_archivo)

    # Abrimos el archivo y lo enviamos
    file_txt = open(nombre_archivo, 'rb')
    
    return file_txt
